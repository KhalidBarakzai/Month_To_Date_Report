#Month-to-Date Misload Report Generator (MTD)
#Author: Khalid Barakzai
#June, 2023

"""
This code generates a month-to-date misloads report based on multiple Word documents in a given folder.
It calculates the misload counts for each loader in the current month and creates a report with the loader names
and their corresponding misload counts. The report includes the date of the report and is saved in a specified output folder.

The code consists of the following functions:

1. get_loader_misloads(doc): Extracts the loader names and misload counts from a Word document's table.

2. calculate_month_to_date_misloads(folder_path): Iterates through the Word documents in the given folder,
   retrieves the loader misloads using the get_loader_misloads function, and calculates the month-to-date misloads
   for each loader.

3. create_month_to_date_report(month_to_date_misloads, output_folder): Generates the month-to-date misloads report
   in a new Word document. It includes the report title, date, a table with the loader names and misload counts,
   and saves the document in the specified output folder.

Usage example:

input_folder = "/path/to/input_folder"
output_folder = "/path/to/output_folder"
month_to_date_misloads = calculate_month_to_date_misloads(input_folder)
create_month_to_date_report(month_to_date_misloads, output_folder)
"""
import os
from docx import Document
from collections import defaultdict
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

##---------------------------------------------------------------------------------------------------

def get_loader_misloads(doc):
    table = doc.tables[0]
    loader_misloads = defaultdict(int)  # Default value for new loaders is 0
    loader_appearances = defaultdict(int)  # Default value for new loaders is 0

    for row in table.rows[1:]:
        loader_name = row.cells[0].text
        misload_count = int(row.cells[1].text)

        if "training" not in loader_name.lower() and "new hire" not in loader_name.lower():
            loader_misloads[loader_name] += misload_count
            loader_appearances[loader_name] += 1

    zero_misload_loaders = table.columns[3].cells[1:]  # Exclude the header cell

    for cell in zero_misload_loaders:
        loader_name = cell.text.strip()  # Remove leading/trailing whitespaces
        if loader_name != "" and "training" not in loader_name.lower() and "new hire" not in loader_name.lower():
            loader_misloads[loader_name] = 0
            loader_appearances[loader_name] += 1

    return loader_misloads, loader_appearances

##---------------------------------------------------------------------------------------------------

def calculate_month_to_date_misloads(folder_path):
    current_month = datetime.now().strftime("%Y-%m")
    month_to_date_misloads = defaultdict(int)
    loader_appearances = defaultdict(int)

    for file_name in os.listdir(folder_path):
        if file_name.endswith(".docx") and current_month in file_name:
            file_path = os.path.join(folder_path, file_name)
            doc = Document(file_path)
            loader_misloads, loader_appearances_single = get_loader_misloads(doc)

            for loader_name, misload_count in loader_misloads.items():
                month_to_date_misloads[loader_name] += misload_count
                loader_appearances[loader_name] += loader_appearances_single[loader_name]

    return month_to_date_misloads, loader_appearances

##---------------------------------------------------------------------------------------------------

def calculate_ratio(misload_count, load_count):
    if misload_count == 0:
        return f"0:{load_count}"
    
    ratio = round(misload_count / load_count, 2)

    if misload_count % load_count == 0:
        ratio = int(ratio)

    return f"{ratio}:1"

##---------------------------------------------------------------------------------------------------

def create_month_to_date_report(month_to_date_misloads, loader_appearances, output_folder):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = 12

    title = doc.add_heading("Three Rivers East: Month-to-Date Misload Report", level=1)
    title.bold = True
    title.runs[0].font.size = Pt(15)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add the report date
    report_date = datetime.now().strftime("%Y-%m-%d")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(f"Report Date: {report_date}")
    run.bold = True

    # Get the loaders with their misload counts
    loaders = list(month_to_date_misloads.keys())

    # Sort loaders based on misload counts in ascending order
    sorted_loaders = sorted(loaders, key=lambda x: month_to_date_misloads[x])

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    header_cells = table.rows[0].cells
    header_cells[0].text = "Loader"
    header_cells[0].paragraphs[0].runs[0].bold = True
    header_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    header_cells[1].text = "Month-to-Date Misloads"
    header_cells[1].paragraphs[0].runs[0].bold = True
    header_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    header_cells[2].text = "Days-Loading"
    header_cells[2].paragraphs[0].runs[0].bold = True
    header_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    header_cells[3].text = "Misloads to Days-Loading Ratio"
    header_cells[3].paragraphs[0].runs[0].bold = True
    header_cells[3].paragraphs[0].runs[0].font.size = Pt(12)

    for loader_name in sorted_loaders:
        misload_count = month_to_date_misloads[loader_name]
        load_count = loader_appearances[loader_name]
        ratio = calculate_ratio(misload_count, load_count)
        cells = table.add_row().cells
        cells[0].text = loader_name
        cells[1].text = str(misload_count)
        cells[2].text = str(load_count)
        cells[3].text = ratio

    # Add loaders with 0 misloads
    for loader_name in loaders:
        if loader_name not in sorted_loaders:
            misload_count = month_to_date_misloads[loader_name]
            load_count = loader_appearances[loader_name]
            ratio = calculate_ratio(misload_count, load_count)
            cells = table.add_row().cells
            cells[0].text = loader_name
            cells[1].text = str(misload_count)
            cells[2].text = str(load_count)
            cells[3].text = ratio

    output_file_name = f"Month_to_Date_Report_{datetime.now().date()}.docx"
    output_path = os.path.join(output_folder, output_file_name)
    doc.save(output_path)

    print(f"Month-to-Date report document saved at: {output_path}")

##---------------------------------------------------------------------------------------------------

input_folder = "/home/kbarakzai/Desktop/TR East Misload Reports"
output_folder = "/home/kbarakzai/Desktop/Month_to_Date Reports"
month_to_date_misloads, loader_appearances = calculate_month_to_date_misloads(input_folder)
create_month_to_date_report(month_to_date_misloads, loader_appearances, output_folder)
